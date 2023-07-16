# from docx import Document
# file = 'template.docx'

# doc = Document(file)


# for para in doc.paragraphs:
#     for run in para.runs:
#         print(run.text)

# word_to_replace = 'Lorem ipsum'
# replacement_word='Thank you for your interest in our company. We are pleased to inform you that you have been selected for the position of [Position]. We believe your skills and experience make you an excellent fit for our team.'

# # Iterate through paragraphs and runs
# for paragraph in doc.paragraphs:
#     for run in paragraph.runs:
#         # Replace the word if found
#         print(run.text)
#         if word_to_replace in run.text:
#             run.text = run.text.replace(word_to_replace, replacement_word)

# section = doc.sections[0]

# for section in doc.sections:
#     # Get the header of each section
#     header = section.header

#     # Print the header content
#     print("Header:")
#     for paragraph in header.paragraphs:
#         print(paragraph.text)


# doc.save('done.docx')


from docx import Document

document_path = 'done.docx'  # Path to your Word document

# # Load the document
# doc = Document(document_path)

# # Initialize an empty list to store the words
# words = []

# # Iterate through paragraphs and runs
# for paragraph in doc.paragraphs:
#     for run in paragraph.runs:
#         # Split the run text into words
#         words.extend(run.text.split())

# # Print each word
# for word in words:
#     print(word)

# from docx import Document

# def replace_placeholder(document, placeholder, replacement):
#     for paragraph in document.paragraphs:
#         if placeholder in paragraph.text:
#             inline = paragraph.runs
#             for i in range(len(inline)):
#                 if placeholder in inline[i].text:
#                     text = inline[i].text.replace(placeholder, replacement)
#                     inline[i].text = text

# # Load the Word document
doc = Document('template.docx')

# # Replace the header placeholder
# replace_placeholder(doc, 'Name', 'Your Header Text')

# # Replace the side heading placeholder
# # replace_placeholder(doc, '<side_heading>', 'Your Side Heading Text')

# # Save the modified document
# doc.save('modified_document.docx')

for paragraph in doc.paragraphs:
    print(paragraph.text)

# Access and print the text in each table
for table in doc.tables:
    #print(table.text)
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    if 'Position' in run.text:
                        text = run.text.replace('Position','Hit')
                        run.text = text
            #print(cell.text)
print(doc.sections[0].header.paragraphs[0].text)
doc.save('now_check.docx')